
from datetime import datetime
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)   # dark blue
LIGHT_GREY = "D9D9D9"


def _add_field(paragraph, field_code: str, placeholder_text: str = ""):
    """Insert a Word field code (e.g. TOC, PAGE) into a paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    text_el = OxmlElement("w:t")
    text_el.text = placeholder_text

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(text_el)
    r.append(fld_end)


def _shade_cell(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_base_styles(document: Document):
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for level, size, color in [(1, 18, ACCENT_COLOR), (2, 14, ACCENT_COLOR)]:
        h_style = document.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        h_style.font.size = Pt(size)
        h_style.font.color.rgb = color
        h_style.font.bold = True


def _add_footer_page_number(document: Document):
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_field(p, "PAGE", "1")


def _add_title_page(document: Document, title: str, document_type: str, audience: str):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    for _ in range(4):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(document_type)
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    document.add_paragraph()

    meta_lines = [
        f"Prepared for: {audience}",
        f"Prepared by: Autonomous AI Agent",
        f"Date: {datetime.now().strftime('%B %d, %Y')}",
    ]
    for line in meta_lines:
        mp = document.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mrun = mp.add_run(line)
        mrun.font.size = Pt(11)
        mrun.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    document.add_page_break()


def _add_toc(document: Document):
    heading = document.add_heading("Table of Contents", level=1)
    p = document.add_paragraph()
    _add_field(p, 'TOC \\o "1-2" \\h \\z \\u', "Right-click here and choose \"Update Field\" to generate the table of contents.")
    note = document.add_paragraph()
    note_run = note.add_run("(Field updates automatically when opened/printed in Microsoft Word, or via right-click \u2192 Update Field.)")
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    document.add_page_break()


def _add_table(document: Document, table_data: Dict[str, Any]):
    headers: List[str] = table_data.get("headers", [])
    rows: List[List[str]] = table_data.get("rows", [])
    if not headers or not rows:
        return

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        _shade_cell(hdr_cells[i], LIGHT_GREY)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    document.add_paragraph()


def _add_section(document: Document, section: Dict[str, Any]):
    document.add_heading(section.get("heading", "Section"), level=1)

    for para_text in section.get("paragraphs", []) or []:
        if para_text:
            document.add_paragraph(para_text)

    bullets = section.get("bullets") or []
    for b in bullets:
        document.add_paragraph(b, style="List Bullet")

    table = section.get("table")
    if table:
        _add_table(document, table)

    document.add_paragraph()


def build_document(
    *,
    title: str,
    document_type: str,
    audience: str,
    sections: List[Dict[str, Any]],
    output_path: str,
) -> str:
    document = Document()
    _set_base_styles(document)
    _add_footer_page_number(document)
    _add_title_page(document, title, document_type, audience)
    _add_toc(document)

    for section in sections:
        _add_section(document, section)

    document.save(output_path)
    return output_path
